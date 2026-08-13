import io

from docx import Document

from app.models.candidate import Candidate
from app.models.job import Job


def build_cv_analysis_docx(candidate: Candidate, job: Job | None, analysis: dict) -> io.BytesIO:
    document = Document()

    document.add_heading("CV Analiz Raporu", level=0)
    document.add_paragraph(f"Aday: {candidate.full_name}")
    document.add_paragraph(f"Pozisyon: {job.title if job else 'Belirtilmemiş'}")

    document.add_heading("Güçlü Yönler", level=1)
    for item in analysis.get("strengths") or []:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Gelişim Alanları", level=1)
    for item in analysis.get("weaknesses") or []:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Özet", level=1)
    document.add_paragraph(analysis.get("summary") or "")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
